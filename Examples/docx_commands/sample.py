from docx import Document
from docx.shared import Inches
import logging

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')

document.save('demo.docx')